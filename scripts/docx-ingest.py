#!/usr/bin/env python3
"""
docx-ingest.py  —  ETSI DOCX  →  corpus/specs/<stem>.json

Produces the same corpus schema as pdf-ingest.py.
Called by ingest.py when a .format-comparison.json recommends 'docx',
or directly for a single file.

Key advantage over PDF: python-docx reads Word Open XML directly,
so tables are extracted as structured rows — no positional heuristics.

Usage:
    python scripts/docx-ingest.py path/to/file.docx
    python scripts/docx-ingest.py downloads/specs/TS/          # batch
    python scripts/docx-ingest.py path/to/file.docx --force    # overwrite
"""

import argparse
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

try:
    import docx as python_docx
except ImportError:
    sys.exit("[ERROR] python-docx not found. Please: pip install python-docx")

# Re-use shortname helpers from pdf-ingest via direct import when run standalone,
# or receive them as arguments when called from ingest.py.
# Fallback: inline the same logic so this file stays self-contained.
try:
    sys.path.insert(0, str(Path(__file__).parent))
    from pdf_ingest import get_shortname, parse_filename as _parse_filename_pdf, SKIP_PREFIXES
except ImportError:
    # Standalone fallback — minimal version
    SKIP_PREFIXES = (".",)

    def get_shortname(norm: str, titles_dir=None):
        return "", ""

    def _parse_filename_pdf(p):
        return p.stem, "vX.X.X"


# ─── Filename parsing (reuse PDF logic, same naming convention) ──────────────

def parse_filename(docx_path: Path) -> tuple[str, str]:
    """Reuse pdf-ingest filename parser — ETSI uses the same stem for both formats."""
    # Temporarily rename stem as if .pdf so the regex matches
    fake_pdf = docx_path.with_suffix(".pdf")
    return _parse_filename_pdf(fake_pdf)


# ─── DOCX text + table extraction ───────────────────────────────────────────

SECTION_RE = re.compile(
    r"^(?P<num>[A-Z]?\.?(?:\d+\.)+\d*|\.?\d+)\s+(?P<title>[A-Z][^\n]{2,80})$",
    re.MULTILINE,
)


def table_to_text(table) -> str:
    """Render a docx Table as pipe-delimited plain text."""
    lines = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        # Deduplicate merged cells (python-docx repeats merged cell content)
        deduped = [cells[0]]
        for c in cells[1:]:
            if c != deduped[-1]:
                deduped.append(c)
        lines.append(" | ".join(deduped))
    return "\n".join(lines)


def extract_document_text(doc_path: Path) -> tuple[list[dict], int]:
    """
    Extract content from a DOCX file as a list of logical "pages".

    DOCX has no concept of pages, so we simulate page boundaries by grouping
    paragraphs into chunks of ~3000 characters each (roughly one A4 page of text).
    Tables are rendered inline as pipe-delimited text.

    Returns (pages, estimated_page_count).
    """
    doc = python_docx.Document(str(doc_path))

    # Collect all content items in document order
    # We iterate doc.element.body to preserve table/paragraph interleaving
    from docx.oxml.ns import qn
    body = doc.element.body

    # Build paragraph-and-table index for lookup
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    PAGE_CHARS = 3000  # ~1 A4 page

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = para_map.get(child)
            if para is None:
                continue
            text = para.text.strip()
            if not text:
                continue
            current.append(text)
            current_len += len(text)
        elif tag == "tbl":
            tbl = table_map.get(child)
            if tbl is None:
                continue
            tbl_text = table_to_text(tbl)
            current.append(tbl_text)
            current_len += len(tbl_text)

        if current_len >= PAGE_CHARS:
            chunks.append("\n".join(current))
            current = []
            current_len = 0

    if current:
        chunks.append("\n".join(current))

    page_count = max(1, len(chunks))
    pages: list[dict] = []
    current_section = ""
    current_title = ""

    for i, chunk in enumerate(chunks, 1):
        hits = [(m.group("num").rstrip("."), m.group("title").strip())
                for m in SECTION_RE.finditer(chunk)]
        if hits:
            current_section, current_title = hits[-1]
        pages.append({
            "page_nr":       i,
            "anchor":        f"#section={current_section}" if current_section else f"#chunk={i}",
            "section":       current_section,
            "section_title": current_title,
            "text_clean":    chunk.strip(),
        })

    return pages, page_count


# ─── Output path ─────────────────────────────────────────────────────────────

def output_path(docx_path: Path, corpus_root: Path) -> Path:
    specs_dir = corpus_root / "specs"
    specs_dir.mkdir(parents=True, exist_ok=True)
    # Use the same stem as the DOCX so it overwrites the PDF-derived JSON
    return specs_dir / (docx_path.stem + ".json")


# ─── Core ingest ─────────────────────────────────────────────────────────────

def ingest_docx(
    docx_path: Path,
    corpus_root: Path,
    titles_dir: Path | None = None,
    verbose: bool = True,
) -> dict:
    norm, version        = parse_filename(docx_path)
    shortname, sn_source = get_shortname(norm, titles_dir)
    shortname_label      = f"  [{shortname}]" if shortname else ""

    if verbose:
        src_label = f" ({sn_source})" if sn_source else ""
        print(f"[INGEST-DOCX] {docx_path.name}  →  {norm} {version}{shortname_label}{src_label}")

    pages, page_count = extract_document_text(docx_path)

    doc = {
        "norm":            norm,
        "shortname":       shortname,
        "shortnameSource": sn_source,
        "shortTitleAI":    None,
        "version":         version,
        "source_pdf":      str(docx_path),   # field kept for schema compat; may be .docx
        "source_format":   "docx",
        "ingested_at":     datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "page_count":      page_count,
        "pages":           pages,
    }

    out = output_path(docx_path, corpus_root)
    out.write_text(json.dumps(doc, ensure_ascii=False, indent=2), encoding="utf-8")

    if verbose:
        chars = sum(len(p["text_clean"]) for p in pages)
        print(f"[OK]          {out}  ({page_count} chunks ≈ pages, {chars:,} Zeichen)")
    return doc


# ─── Batch helpers ────────────────────────────────────────────────────────────

def collect_docx(inputs: list[str]) -> list[Path]:
    result: list[Path] = []
    for inp in inputs:
        p = Path(inp)
        if p.is_dir():
            for f in sorted(p.rglob("*.docx")):
                if not f.name.startswith(SKIP_PREFIXES):
                    result.append(f)
        elif p.is_file() and p.suffix.lower() in (".docx", ".doc"):
            if not p.name.startswith(SKIP_PREFIXES):
                result.append(p)
        else:
            print(f"[WARN] {inp} is not a .docx file or directory — skipped.", file=sys.stderr)
    return result


# ─── CLI ──────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="ETSI DOCX → corpus/specs/*.json  (idempotent by default)",
    )
    parser.add_argument("input", nargs="+", help="DOCX file(s) or directories")
    parser.add_argument("--corpus", default="corpus")
    parser.add_argument("--titles", default="downloads/specs/_titles")
    parser.add_argument("--force", "-f", action="store_true")
    parser.add_argument("--quiet", "-q", action="store_true")
    args = parser.parse_args()

    corpus_root = Path(args.corpus)
    titles_dir  = Path(args.titles)
    verbose     = not args.quiet
    files       = collect_docx(args.input)

    if not files:
        sys.exit("[ERROR] No .docx files found.")

    if verbose:
        mode = "--force: always overwrite" if args.force else "idempotent: skip existing"
        print(f"[INFO] {len(files)} DOCX files found  ({mode})")

    ok = skipped = errors = 0
    for f in files:
        out = output_path(f, corpus_root)
        if not args.force and out.exists():
            if verbose:
                print(f"[SKIP] {f.name}")
            skipped += 1
            continue
        try:
            ingest_docx(f, corpus_root, titles_dir=titles_dir, verbose=verbose)
            ok += 1
        except Exception as exc:
            print(f"[ERROR] {f.name}: {exc}", file=sys.stderr)
            errors += 1

    if verbose:
        print(f"\n[DONE] {ok} processed | {skipped} skipped | {errors} errors  (of {len(files)} total)")


if __name__ == "__main__":
    main()
