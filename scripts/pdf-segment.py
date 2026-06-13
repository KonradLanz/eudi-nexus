#!/usr/bin/env python3
"""
pdf-segment.py  —  ETSI document structural segmentation pipeline

Primary path  : DOCX  (python-docx — structured paragraphs + real tables)
Fallback path : PDF   (pdfplumber — geometry-based, no auth needed)

Second-pass over corpus/specs/*.json (produced by pdf-ingest.py / docx-ingest.py).
For each page/paragraph the raw text is re-analysed to separate:

  HEADER   — repeating page header  (doc title, spec number, ETSI logo line)
  FOOTER   — page footer            (page number, date, copyright)
  TOC      — table of contents
  SECTION  — heading / clause title
  NORM     — normative body text    (RFC-2119: shall/shall not/…)
  INFORM   — informative text       (NOTE, Example, Annex (informative), …)
  TABLE    — table / figure caption
  OTHER    — cover page, boilerplate, blank

Output per spec
  corpus/specs/_segments/<stem>.segments.json   — machine-readable segments
  corpus/specs/_adoc/<stem>.adoc                — AsciiDoc with [[anchor]] backrefs

Style detection
  ETSI DOCX files use a consistent Word template across all specs.
  See docs/etsi-style-map.md for the full style reference and AI fallback guidance.

Fair-use note:
  Local AI processing of standards documents for compliance research
  constitutes fair use under EU and international copyright law
  (Art. 5(3) InfoSoc Directive). No redistribution of document content.
  Auth credentials are never stored in this repository.

Usage:
  python3 scripts/pdf-segment.py                    # all corpus specs (DOCX preferred)
  python3 scripts/pdf-segment.py corpus/specs/ts_119403v020201p.json
  python3 scripts/pdf-segment.py --pdf path/to/doc.pdf
  python3 scripts/pdf-segment.py --docx path/to/doc.docx
  python3 scripts/pdf-segment.py --force             # overwrite existing
  python3 scripts/pdf-segment.py --profile etsi-contribution --pdf ESI.pdf
  python3 scripts/pdf-segment.py --scan-styles path/to/doc.docx  # style audit only

Dependencies:
  pip install pdfplumber python-docx
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import docx as docx_lib
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ─────────────────────────────────────────────────────────────────────────────
# Magic-byte file-type detection
# ─────────────────────────────────────────────────────────────────────────────

_ZIP_MAGIC = b"PK\x03\x04"
_PDF_MAGIC = b"%PDF"


def _sniff_kind(path: Path) -> str:
    """
    Read the first 4 bytes of *path* and return 'docx', 'pdf', or 'unknown'.

    ETSI sometimes serves DOCX files under a .pdf URL / filename.
    This check catches that mislabelling before pdfplumber tries to open the
    file and raises "No /Root object! - Is this really a PDF?".
    """
    try:
        header = path.read_bytes()[:4]
    except OSError:
        return "unknown"
    if header.startswith(_PDF_MAGIC):
        return "pdf"
    if header.startswith(_ZIP_MAGIC):
        return "docx"   # ZIP → treat as DOCX (Office Open XML)
    return "unknown"


# ─────────────────────────────────────────────────────────────────────────────
# Profiles
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class Profile:
    name: str
    header_zone:        float = 0.10
    footer_zone:        float = 0.08
    heading_min_size:   float = 11.5
    heading_bold_ratio: float = 0.55
    toc_run_threshold:  int   = 4
    header_re:      list[str] = field(default_factory=list)
    footer_re:      list[str] = field(default_factory=list)
    boilerplate_re: list[str] = field(default_factory=list)


PROFILES: dict[str, Profile] = {
    "etsi-spec": Profile(
        name="etsi-spec",
        header_zone=0.10, footer_zone=0.09,
        heading_min_size=11.0, heading_bold_ratio=0.50,
        toc_run_threshold=4,
        header_re=[
            r"ETSI\s+(EN|TS|TR|GS|GR|EG)\s+\d",
            r"Draft\s+ETSI",
            r"^ETSI$",
        ],
        footer_re=[
            r"^\d+$",
            r"ETSI\s*$",
            r"\d{4}-\d{2}$",
            r"\u00a9\s*ETSI",
            r"Publicly Available",
            r"Restricted",
            r"Confidential",
        ],
        boilerplate_re=[
            r"^Reference\s*$",
            r"^Keywords\s*$",
            r"^ETSI\s+650 Route des Lucioles",
            r"Important Notice",
            r"Intellectual Property Rights",
            r"Essential patents",
            r"No guarantee can be given",
            r"Modal verbs terminology",
            r"In the present document",
        ],
    ),
    "etsi-contribution": Profile(
        name="etsi-contribution",
        header_zone=0.12, footer_zone=0.10,
        heading_min_size=11.0, heading_bold_ratio=0.50,
        toc_run_threshold=3,
        header_re=[
            r"ETSI TC ESI",
            r"ESI\(\d{2}\)\d+",
            r"Source:\s+",
            r"^Meeting\b",
        ],
        footer_re=[
            r"^\d+$",
            r"\u00a9\s*ETSI",
            r"Publicly Available",
            r"Confidential",
        ],
        boilerplate_re=[
            r"The present document has been",
        ],
    ),
    "ietf-rfc": Profile(
        name="ietf-rfc",
        header_zone=0.06, footer_zone=0.06,
        heading_min_size=10.5, heading_bold_ratio=0.40,
        toc_run_threshold=4,
        header_re=[
            r"^RFC\s+\d+",
            r"Internet Engineering Task Force",
            r"^\[?RFC\d+\]?",
        ],
        footer_re=[
            r"^\[Page \d+\]$",
            r"Internet Standards Track",
        ],
        boilerplate_re=[
            r"Status of This Memo",
            r"Copyright Notice",
            r"ISSN:",
        ],
    ),
}


# ─────────────────────────────────────────────────────────────────────────────
# Shared regexes
# ─────────────────────────────────────────────────────────────────────────────

_RFC2119_RE = re.compile(
    r"\b(shall\s+not|shall|should\s+not|should|must\s+not|must"
    r"|required|recommended|may\s+not|may|optional)\b",
    re.IGNORECASE,
)
_INFORMATIVE_RE = re.compile(
    r"^\s*(NOTE\b|EXAMPLE\b|Example\s+\d|\[i\.\d+\])"
    r"|Annex\s+\w+\s+\(informative\)"
    r"|\(informative\)",
    re.IGNORECASE | re.MULTILINE,
)
_TOC_LINE_RE = re.compile(
    r"^.{2,60}[\.\ ]{4,}\d{1,4}\s*$",
    re.MULTILINE,
)
_SECTION_RE = re.compile(
    r"^(?P<num>[A-Z]?\.?(?:\d+\.)+\d*|\.?\d+)\s+(?P<title>[A-Z][^\n]{2,80})$",
    re.MULTILINE,
)
_BOILERPLATE_RE = re.compile(
    r"Intellectual Property|Essential patents|Important Notice"
    r"|Modal verbs terminology|Foreword",
    re.IGNORECASE,
)

# Matches real table/figure captions — NOT NOTE/EXAMPLE prose blocks.
# NOTE and EXAMPLE are informative text; they must be caught by
# _INFORMATIVE_RE → INFORM, not classified as TABLE.
_FIGURE_TABLE_CAPTION_RE = re.compile(
    r"^(Figure|Table)\s+[\d\-A-Z]",
    re.IGNORECASE,
)

# Sentence-ending punctuation — used by the paragraph-merge heuristic
_SENTENCE_END_RE = re.compile(r"[.!?:]\s*$")

SEGMENT_TYPES = (
    "HEADER", "FOOTER", "TOC", "SECTION",
    "NORM", "INFORM", "TABLE", "OTHER",
)

# ── ETSI Word template — known style → segment type mapping ──────────────────
_ETSI_STYLE_MAP: dict[str, str] = {
    # ── Cover page ────────────────────────────────────────────────
    "za":           "OTHER",
    "zb":           "OTHER",
    "zt":           "OTHER",
    "zd":           "OTHER",
    "fp":           "OTHER",

    # ── Informative ───────────────────────────────────────────────
    "no":           "INFORM",
    "ex":           "INFORM",
    "ew":           "INFORM",
    "editor's note": "INFORM",

    # ── Normative body text ───────────────────────────────────────
    "b1":           "NORM",
    "b1+":          "NORM",
    "b2":           "NORM",
    "b2+":          "NORM",
    "b3":           "NORM",
    "b3+":          "NORM",
    "b4":           "NORM",
    "bl":           "NORM",
    "bn":           "NORM",
    "pl":           "NORM",
    "list paragraph": "NORM",

    # ── Table cells ───────────────────────────────────────────────
    "th":           "TABLE",
    "tf":           "TABLE",

    # ── Annex headings ────────────────────────────────────────────
    "h6":           "SECTION",
}

_KNOWN_STYLE_PREFIXES = (
    "heading",
    "toc ",
    "normal",
    "default paragraph",
    "za", "zb", "zt", "zd", "fp",
    "no", "ex", "ew", "editor",
    "b1", "b1+", "b2", "b2+", "b3", "b3+", "b4",
    "bl", "bn", "pl", "list paragraph",
    "th", "tf",
    "h6",
    "tt",
)


def find_normative_keywords(text: str) -> list[str]:
    return list({m.group(1).lower() for m in _RFC2119_RE.finditer(text)})


def _safe_stem(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")


# ─────────────────────────────────────────────────────────────────────────────
# Dynamic style detection  (scan once per DOCX, cache result)
# ─────────────────────────────────────────────────────────────────────────────

_style_cache: dict[str, set[str]] = {}


def detect_docx_heading_styles(docx_path: Path) -> tuple[set[str], set[str]]:
    key = str(docx_path)
    if key in _style_cache:
        cached = _style_cache[key]
        known   = {s for s in cached if _is_known_style(s)}
        unknown = cached - known
        return known, unknown

    doc = docx_lib.Document(str(docx_path))
    all_styles: set[str] = set()
    for p in doc.paragraphs:
        if p.text.strip():
            all_styles.add((p.style.name or "").lower())

    _style_cache[key] = all_styles
    known   = {s for s in all_styles if _is_known_style(s)}
    unknown = all_styles - known
    return known, unknown


def _is_known_style(style_lower: str) -> bool:
    return any(style_lower.startswith(p) for p in _KNOWN_STYLE_PREFIXES)


def _warn_unknown_styles(unknown: set[str], docx_path: Path) -> None:
    if not unknown:
        return
    flagged = {s for s in unknown if len(s) > 1 and not s.isdigit()}
    if not flagged:
        return
    print(
        f"   ⚠️  Unknown styles in {docx_path.name} "
        f"(may need AI classification): "
        f"{', '.join(sorted(flagged)[:8])}",
        file=sys.stderr,
    )


# ─────────────────────────────────────────────────────────────────────────────
# DOCX segmentation  (primary path — python-docx)
# ─────────────────────────────────────────────────────────────────────────────

def _docx_para_is_bold(para) -> bool:
    runs = para.runs
    if not runs:
        return False
    bold_runs = sum(1 for r in runs if r.bold)
    return bold_runs / len(runs) >= 0.5


def _docx_para_type(para, profile: Profile) -> str:
    style_raw  = para.style.name or ""
    style_low  = style_raw.lower()
    text       = para.text.strip()

    if not text:
        return "OTHER"

    mapped = _ETSI_STYLE_MAP.get(style_low)
    if mapped:
        return mapped

    if re.match(r"heading", style_low, re.IGNORECASE):
        return "SECTION"

    if re.match(r"toc\s", style_low) or "contents" in style_low:
        return "TOC"

    if _BOILERPLATE_RE.search(text):
        return "OTHER"

    if _TOC_LINE_RE.search(text):
        return "TOC"

    if _INFORMATIVE_RE.search(text):
        return "INFORM"

    if _docx_para_is_bold(para) and len(text) < 120:
        m = _SECTION_RE.match(text)
        if m:
            return "SECTION"

    if _FIGURE_TABLE_CAPTION_RE.match(text):
        return "TABLE"

    if find_normative_keywords(text):
        return "NORM"

    return "OTHER"


def segment_docx(
    docx_path: Path,
    profile: Profile,
    doc_stem: str | None = None,
) -> list[dict]:
    if not HAS_DOCX:
        raise RuntimeError(
            "python-docx not installed — run: pip install python-docx"
        )

    stem = doc_stem or _safe_stem(docx_path.stem)
    doc  = docx_lib.Document(str(docx_path))

    _, unknown_styles = detect_docx_heading_styles(docx_path)
    _warn_unknown_styles(unknown_styles, docx_path)

    segments: list[dict] = []
    seg_counters: dict[int, int] = {}
    current_section       = ""
    current_section_title = ""
    para_idx = 0

    def _new_seg(seg_type: str, text: str, para_num: int) -> dict:
        seg_counters[para_num] = seg_counters.get(para_num, 0) + 1
        seg_id = f"{stem}_pa{para_num}_b{seg_counters[para_num]}"
        kw = find_normative_keywords(text) if seg_type in ("NORM", "SECTION", "OTHER") else []
        return {
            "id":                  seg_id,
            "type":               seg_type,
            "page":               para_num,
            "anchor":             f"#para={para_num}",
            "section":            current_section,
            "section_title":      current_section_title,
            "text":               text,
            "normative_keywords": kw,
            "profile":            profile.name,
        }

    body      = doc.element.body
    para_map  = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = para_map.get(child)
            if para is None:
                continue
            text     = para.text.strip()
            para_idx += 1
            if not text:
                continue

            seg_type = _docx_para_type(para, profile)

            if seg_type == "SECTION":
                m = _SECTION_RE.match(text)
                if m:
                    current_section       = m.group("num").rstrip(".")
                    current_section_title = m.group("title").strip()
                else:
                    current_section_title = text

            segments.append(_new_seg(seg_type, text, para_idx))

        elif tag == "tbl":
            tbl = table_map.get(child)
            if tbl is None:
                continue
            para_idx += 1

            rows = []
            for row in tbl.rows:
                cells = [cell.text.strip() for cell in row.cells]
                deduped: list[str] = []
                for c in cells:
                    if not deduped or c != deduped[-1]:
                        deduped.append(c)
                rows.append(deduped)

            if not rows:
                continue

            header    = rows[0]
            data_rows = rows[1:]
            flat_text = _table_to_markdown(rows)

            seg_counters[para_idx] = seg_counters.get(para_idx, 0) + 1
            seg_id = f"{stem}_pa{para_idx}_tbl{seg_counters[para_idx]}"

            segments.append({
                "id":                  seg_id,
                "type":               "TABLE",
                "page":               para_idx,
                "anchor":             f"#para={para_idx}",
                "section":            current_section,
                "section_title":      current_section_title,
                "text":               flat_text,
                "markdown_text":      flat_text,
                "table_header":       header,
                "table_rows":         data_rows,
                "normative_keywords": find_normative_keywords(flat_text),
                "profile":            profile.name,
            })

    return segments


# ─────────────────────────────────────────────────────────────────────────────
# PDF segmentation  (fallback — pdfplumber)
# ─────────────────────────────────────────────────────────────────────────────

_GAP_THRESHOLD = 20
RULES_PATH = Path("corpus/para-tune/learned-rules.json")


def _get_rule_for_stem(stem: str) -> dict:
    """
    Load learned para-tune overrides directly from JSON.

    Returns an empty dict if the rules file is missing, invalid, or the stem
    has no entry. A shallow copy is returned so callers/tests cannot mutate the
    cached on-disk data through the returned object.
    """
    if not RULES_PATH.exists():
        return {}
    try:
        rules = json.loads(RULES_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}
    rule = rules.get(stem, {})
    return dict(rule) if isinstance(rule, dict) else {}


# Fraction of page height — tables touching these zones are candidates for
# cross-page continuation detection.
_PAGE_BOTTOM_ZONE = 0.85   # table ending below this y-fraction → may continue
_PAGE_TOP_ZONE    = 0.15   # table starting above this y-fraction on next page

# ── Table quality guards ──────────────────────────────────────────────────────
# ETSI PDFs often render NOTE/EXAMPLE blocks inside invisible single-cell
# frames.  pdfplumber's find_tables() detects these as 1×1 tables.
# These constants define the minimum size a detected table must have to be
# accepted as a real data table (not a framed prose block).
_MIN_TABLE_COLS     = 2     # at least 2 columns
_MIN_TABLE_ROWS     = 2     # at least 1 header row + 1 data row
_MIN_TABLE_TEXT_LEN = 80    # total cell text must be longer than a single NOTE


@dataclass
class TextBlock:
    text:          str
    y_top:         float
    y_bot:         float
    page_height:   float
    avg_font_size: float
    bold_ratio:    float

    @property
    def y_frac_top(self) -> float:
        return 1.0 - (self.y_bot / self.page_height)

    @property
    def y_frac_bot(self) -> float:
        return 1.0 - (self.y_top / self.page_height)


# ─────────────────────────────────────────────────────────────────────────────
# Table helpers
# ─────────────────────────────────────────────────────────────────────────────

def _table_to_markdown(rows: list[list[str]]) -> str:
    """
    Convert a list-of-lists table (as returned by pdfplumber or python-docx)
    into a GitHub-Flavored Markdown table string.

    The first row is treated as the header.  Empty rows are skipped.
    Cell text is stripped and pipe characters inside cells are escaped.
    """
    if not rows:
        return ""

    def _cell(s: str) -> str:
        return s.replace("|", "\\|").strip()

    header = rows[0]
    n_cols = len(header)
    lines: list[str] = []

    # Header row
    lines.append("| " + " | ".join(_cell(c) for c in header) + " |")
    # Separator
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    # Data rows — skip completely empty rows, pad/trim to n_cols
    for row in rows[1:]:
        padded = (list(row) + [""] * n_cols)[:n_cols]
        if all(not c.strip() for c in padded):
            continue
        lines.append("| " + " | ".join(_cell(c) for c in padded) + " |")

    return "\n".join(lines)


def _extract_tables(page) -> list[dict]:
    """
    Extract all *real* tables from a pdfplumber page object.

    Returns a list of dicts, each with:
      bbox          — (x0, top, x1, bottom) in page coordinates
      rows          — list[list[str]]  (raw cell text, including header)
      header        — rows[0]
      data_rows     — rows[1:]
      markdown_text — GFM Markdown string
      y_frac_top    — normalised vertical position (0=top, 1=bottom)
      y_frac_bot    — normalised vertical position

    Quality guards (constants at top of section):
      _MIN_TABLE_COLS     — rejects single-column NOTE/EXAMPLE frames
      _MIN_TABLE_ROWS     — rejects single-row caption boxes
      _MIN_TABLE_TEXT_LEN — rejects very short framed text blocks
    """
    height = float(page.height)
    result = []
    try:
        tables = page.find_tables()
    except Exception:
        return result

    for tbl in tables:
        try:
            raw = tbl.extract()
        except Exception:
            continue
        if not raw:
            continue

        # Normalise: replace None with "", strip whitespace, drop empty rows
        rows: list[list[str]] = []
        for raw_row in raw:
            cleaned = [(c or "").strip() for c in raw_row]
            if any(cleaned):
                rows.append(cleaned)

        if not rows:
            continue

        # ── Quality guard 1: minimum column count ─────────────────
        n_cols = max(len(r) for r in rows)
        if n_cols < _MIN_TABLE_COLS:
            continue

        # ── Quality guard 2: minimum row count ────────────────────
        if len(rows) < _MIN_TABLE_ROWS:
            continue

        # ── Quality guard 3: total text length ────────────────────
        total_text = " ".join(c for row in rows for c in row)
        if len(total_text) < _MIN_TABLE_TEXT_LEN:
            continue

        bbox = tbl.bbox  # (x0, top, x1, bottom)
        y_top = float(bbox[1])
        y_bot = float(bbox[3])

        result.append({
            "bbox":          bbox,
            "rows":          rows,
            "header":        rows[0],
            "data_rows":     rows[1:],
            "markdown_text": _table_to_markdown(rows),
            "y_frac_top":    y_top / height,
            "y_frac_bot":    y_bot / height,
        })

    return result


def _is_table_bbox(y_top: float, y_bot: float, table_bboxes: list[tuple]) -> bool:
    """
    Return True if the vertical range [y_top, y_bot] overlaps with any
    known table bounding box on this page.  Used to mask table regions
    from the text-block extractor so we don't double-count table content.
    """
    for (_, t_top, _, t_bot) in table_bboxes:
        overlap_top = max(y_top, t_top)
        overlap_bot = min(y_bot, t_bot)
        if overlap_bot - overlap_top > 2:   # >2pt overlap
            return True
    return False


def _merge_cross_page_tables(segments: list[dict]) -> list[dict]:
    """
    Detect and merge tables that span a page boundary.

    Strategy A — header repeat (Word "Repeat Header Rows"):
      If page N ends with a TABLE whose last row's y_frac_bot > _PAGE_BOTTOM_ZONE
      AND page N+1 starts with a TABLE whose y_frac_top < _PAGE_TOP_ZONE
      AND both tables have the same number of columns
      AND the first row of the page-N+1 table is identical to the header of
          the page-N table → it is a repeated header; strip it and append rows.

    Strategy B — column-count match only (no repeated header):
      Same positional criteria, same column count, but headers differ →
      set continues_on_next_page=True on the page-N segment so the retrieval
      pipeline can include both chunks together.

    Operates in-place on *segments* (sorted by page then position).
    Returns the (possibly shorter) merged list.
    """
    if not segments:
        return segments

    # Index TABLE segments by page
    table_segs: list[int] = [
        i for i, s in enumerate(segments)
        if s.get("type") == "TABLE" and "table_header" in s
    ]

    merged_indices: set[int] = set()

    for idx in range(len(table_segs) - 1):
        i = table_segs[idx]
        j = table_segs[idx + 1]

        seg_a = segments[i]
        seg_b = segments[j]

        page_a = seg_a.get("page", 0)
        page_b = seg_b.get("page", 0)
        if page_b != page_a + 1:
            continue

        # Positional check
        y_bot_a   = seg_a.get("y_frac_bot", 0.0)
        y_top_b   = seg_b.get("y_frac_top", 1.0)
        if y_bot_a < _PAGE_BOTTOM_ZONE or y_top_b > _PAGE_TOP_ZONE:
            continue

        header_a   = seg_a.get("table_header", [])
        header_b   = seg_b.get("table_header", [])
        data_rows_b = seg_b.get("table_rows", [])

        n_cols_a = len(header_a)
        n_cols_b = len(header_b)
        if n_cols_a != n_cols_b or n_cols_a == 0:
            # Strategy B: just annotate
            seg_a["continues_on_next_page"] = True
            continue

        # Strategy A: repeated header?
        if header_b == header_a and data_rows_b:
            # Merge: append page-B data rows into page-A segment, drop page-B
            merged_rows = seg_a.get("table_rows", []) + data_rows_b
            all_rows    = [header_a] + merged_rows
            seg_a["table_rows"]    = merged_rows
            seg_a["text"]          = _table_to_markdown(all_rows)
            seg_a["markdown_text"] = seg_a["text"]
            seg_a["merged_from_page"] = page_b
            merged_indices.add(j)
        else:
            # Strategy B: column counts match but headers differ — annotate only
            seg_a["continues_on_next_page"] = True

    return [s for idx, s in enumerate(segments) if idx not in merged_indices]


def _extract_blocks(page, table_bboxes: list[tuple] | None = None) -> list[TextBlock]:
    """
    Extract text blocks from a pdfplumber page, skipping characters that fall
    inside known table bounding boxes (so table content isn't double-counted).
    """
    chars = page.chars
    if not chars:
        return []
    height = float(page.height)
    mask   = table_bboxes or []

    lines: dict[int, list[dict]] = {}
    for ch in chars:
        # Skip characters inside table regions
        if mask and _is_table_bbox(ch["top"], ch["bottom"], mask):
            continue
        y_mid = int((ch["top"] + ch["bottom"]) / 2)
        lines.setdefault(y_mid, []).append(ch)

    sorted_ys = sorted(lines)
    blocks: list[TextBlock] = []
    current_chars: list[dict] = []
    current_ys:    list[int]  = []

    def _flush(ys: list[int], chs: list[dict]) -> None:
        if not chs:
            return
        text = "".join(c["text"] for c in sorted(chs, key=lambda c: (c["top"], c["x0"])))
        text = re.sub(r" {2,}", " ", text).strip()
        if not text:
            return
        sizes = [c.get("size", 0) for c in chs if c.get("size")]
        avg_size   = sum(sizes) / len(sizes) if sizes else 0.0
        bold_count = sum(
            1 for c in chs
            if "Bold" in (c.get("fontname") or "") or "bold" in (c.get("fontname") or "")
        )
        bold_ratio = bold_count / len(chs) if chs else 0.0
        y_top = float(min(c["top"]    for c in chs))
        y_bot = float(max(c["bottom"] for c in chs))
        blocks.append(TextBlock(text=text, y_top=y_top, y_bot=y_bot,
                                page_height=height, avg_font_size=avg_size,
                                bold_ratio=bold_ratio))

    for y in sorted_ys:
        if current_ys and (y - current_ys[-1]) > _GAP_THRESHOLD:
            _flush(current_ys, current_chars)
            current_chars = []
            current_ys    = []
        current_chars.extend(lines[y])
        current_ys.append(y)
    _flush(current_ys, current_chars)
    return sorted(blocks, key=lambda b: b.y_top)


def _dedup_blocks(blocks: list[TextBlock]) -> list[TextBlock]:
    if not blocks:
        return blocks
    out: list[TextBlock] = [blocks[0]]
    for b in blocks[1:]:
        prev = out[-1]
        if b.text == prev.text and abs(b.y_top - prev.y_top) < 5:
            continue
        out.append(b)
    return out


def _merge_paragraph_blocks(
    blocks: list[TextBlock],
    seg_types: list[str],
) -> tuple[list[TextBlock], list[str]]:
    if not blocks:
        return blocks, seg_types

    _MERGEABLE_TYPES = {"NORM", "INFORM"}
    _NEW_ITEM_RE = re.compile(
        r"^(\d+[\.\ ]\s"
        r"|\([a-z]\)\s"
        r"|[-–•]\s"
        r"|NOTE\b"
        r"|EXAMPLE\b)",
        re.IGNORECASE,
    )

    merged_blocks: list[TextBlock] = []
    merged_types:  list[str]       = []

    for blk, stype in zip(blocks, seg_types):
        if (
            merged_blocks
            and stype in _MERGEABLE_TYPES
            and merged_types[-1] == stype
            and not _SENTENCE_END_RE.search(merged_blocks[-1].text)
            and not _NEW_ITEM_RE.match(blk.text)
        ):
            prev = merged_blocks[-1]
            merged_text = prev.text.rstrip() + " " + blk.text.lstrip()
            merged_blocks[-1] = TextBlock(
                text=merged_text,
                y_top=prev.y_top,
                y_bot=blk.y_bot,
                page_height=prev.page_height,
                avg_font_size=prev.avg_font_size,
                bold_ratio=prev.bold_ratio,
            )
        else:
            merged_blocks.append(blk)
            merged_types.append(stype)

    return merged_blocks, merged_types


def classify_block(
    block: TextBlock,
    profile: Profile,
    page_nr: int,
    toc_mode: bool,
) -> str:
    text = block.text
    if block.y_frac_top < profile.header_zone:
        if any(re.search(r, text) for r in profile.header_re):
            return "HEADER"
        if page_nr > 1 and block.y_frac_bot < profile.header_zone * 1.5:
            return "HEADER"
    if block.y_frac_bot > (1.0 - profile.footer_zone):
        if any(re.search(r, text) for r in profile.footer_re):
            return "FOOTER"
        if page_nr > 1 and block.y_frac_top > (1.0 - profile.footer_zone * 1.5):
            return "FOOTER"
    if page_nr <= 4 and any(re.search(r, text) for r in profile.boilerplate_re):
        return "OTHER"
    toc_matches = len(_TOC_LINE_RE.findall(text))
    if toc_mode or toc_matches >= profile.toc_run_threshold:
        return "TOC"
    is_heading_font = (
        block.avg_font_size >= profile.heading_min_size
        and block.bold_ratio  >= profile.heading_bold_ratio
    )
    if is_heading_font:
        m = _SECTION_RE.match(text.strip())
        if m:
            return "SECTION"
    # NOTE/EXAMPLE blocks are informative prose — check BEFORE the caption regex
    # so they are never mis-classified as TABLE.
    if _INFORMATIVE_RE.search(text):
        return "INFORM"
    # Only real Figure/Table captions become TABLE segments via the text path.
    # NOTE and EXAMPLE are already handled above.
    if _FIGURE_TABLE_CAPTION_RE.match(text):
        return "TABLE"
    if find_normative_keywords(text):
        return "NORM"
    return "OTHER"


def segment_pdf(
    pdf_path: Path,
    profile: Profile,
    doc_stem: str | None = None,
) -> list[dict]:
    if not HAS_PDFPLUMBER:
        raise RuntimeError("pdfplumber not installed — run: pip install pdfplumber")
    stem = doc_stem or _safe_stem(pdf_path.stem)
    segments: list[dict] = []
    seg_counters: dict[int, int] = {}
    current_section       = ""
    current_section_title = ""
    toc_page_count = 0
    MAX_TOC_PAGES  = 6

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page in pdf.pages:
            page_nr    = page.page_number
            page_h     = float(page.height)

            # ── Pass 0: extract real tables first ─────────────────
            # Tables are extracted with pdfplumber's table finder,
            # then their bboxes are used to mask the text-block extractor
            # so table content isn't double-counted as NORM/INFORM blocks.
            pdf_tables  = _extract_tables(page)
            table_bboxes = [t["bbox"] for t in pdf_tables]

            # Emit TABLE segments (inserted before text blocks on same page
            # so section context is correct for interleaved tables)
            for tbl in pdf_tables:
                seg_counters[page_nr] = seg_counters.get(page_nr, 0) + 1
                seg_id = f"{stem}_p{page_nr}_tbl{seg_counters[page_nr]}"
                segments.append({
                    "id":                  seg_id,
                    "type":               "TABLE",
                    "page":               page_nr,
                    "anchor":             f"#page={page_nr}",
                    "section":            current_section,
                    "section_title":      current_section_title,
                    "text":               tbl["markdown_text"],
                    "markdown_text":      tbl["markdown_text"],
                    "table_header":       tbl["header"],
                    "table_rows":         tbl["data_rows"],
                    "y_frac_top":         tbl["y_frac_top"],
                    "y_frac_bot":         tbl["y_frac_bot"],
                    "normative_keywords": find_normative_keywords(tbl["markdown_text"]),
                    "profile":            profile.name,
                })

            # ── Pass 1: text blocks (masked) ───────────────────────
            raw_blocks = _extract_blocks(page, table_bboxes)

            # ── Pass 2: remove duplicate adjacent blocks ──────────
            blocks = _dedup_blocks(raw_blocks)

            # ── Pass 3: classify all blocks on this page ──────────
            toc_lines = sum(
                1 for b in blocks
                if len(_TOC_LINE_RE.findall(b.text)) >= 1
            )
            in_toc = toc_lines >= profile.toc_run_threshold
            if in_toc:
                toc_page_count += 1
            elif toc_page_count > 0 and toc_page_count < MAX_TOC_PAGES:
                toc_page_count = 0

            toc_mode = toc_page_count > 0 and toc_page_count < MAX_TOC_PAGES
            classified = [
                classify_block(b, profile, page_nr, toc_mode)
                for b in blocks
            ]

            # ── Pass 4: merge split paragraph blocks ──────────────
            blocks, classified = _merge_paragraph_blocks(blocks, classified)

            # ── Emit text segments ────────────────────────────────
            for block, seg_type in zip(blocks, classified):
                if not block.text.strip():
                    continue
                if seg_type == "SECTION":
                    m = _SECTION_RE.match(block.text.strip())
                    if m:
                        current_section       = m.group("num").rstrip(".")
                        current_section_title = m.group("title").strip()
                seg_counters[page_nr] = seg_counters.get(page_nr, 0) + 1
                seg_id = f"{stem}_p{page_nr}_b{seg_counters[page_nr]}"
                kw = find_normative_keywords(block.text) if seg_type in ("NORM", "SECTION", "OTHER") else []
                segments.append({
                    "id":                  seg_id,
                    "type":               seg_type,
                    "page":               page_nr,
                    "anchor":             f"#page={page_nr}",
                    "section":            current_section,
                    "section_title":      current_section_title,
                    "text":               block.text,
                    "normative_keywords": kw,
                    "profile":            profile.name,
                })

    # ── Cross-page table merge (post-processing) ──────────────────
    segments = _merge_cross_page_tables(segments)

    return segments


# ─────────────────────────────────────────────────────────────────────────────
# Source file discovery
# ─────────────────────────────────────────────────────────────────────────────

DOWNLOAD_ROOTS = [
    Path("downloads/specs/TS"),
    Path("downloads/specs/EN"),
    Path("downloads/specs/TR"),
    Path("downloads/specs/SR"),
    Path("downloads/specs"),
]


def find_source_file(rec: dict, stem: str) -> tuple[Path | None, str]:
    docx_str = rec.get("source_docx", "")
    if docx_str:
        p = Path(docx_str)
        if p.is_file():
            return p, "docx"

    pdf_str = rec.get("source_pdf", "")
    if pdf_str:
        p = Path(pdf_str)
        if p.is_file():
            real_kind = _sniff_kind(p)
            if real_kind == "unknown":
                real_kind = "pdf"
            return p, real_kind

    norm_base = re.sub(r"v\d{6}p$", "", stem)
    for root in DOWNLOAD_ROOTS:
        for suffix, kind in [(".docx", "docx"), (".pdf", "pdf")]:
            for p in sorted(root.glob(f"{norm_base}*{suffix}"), reverse=True):
                return p, kind

    return None, ""


# ─────────────────────────────────────────────────────────────────────────────
# Shortname helper
# ─────────────────────────────────────────────────────────────────────────────

def _format_shortname(rec: dict) -> str:
    norm    = rec.get("norm",    "")
    version = rec.get("version", "")
    title   = rec.get("title",   "")
    if not title:
        title = rec.get("doc_title", rec.get("description", ""))
    if not norm:
        norm = rec.get("spec", rec.get("number", ""))
    ver_str   = f" v{version.lstrip('Vv')}" if version else ""
    title_str = f" — {title[:70]}" if title else ""
    return f"{norm}{ver_str}{title_str}".strip()


# ─────────────────────────────────────────────────────────────────────────────
# AsciiDoc export
# ─────────────────────────────────────────────────────────────────────────────

_ADOC_SECTION_LEVEL = {1: "=", 2: "==", 3: "===", 4: "====", 5: "====="}


def _section_depth(section_num: str) -> int:
    if not section_num:
        return 1
    parts = section_num.strip(".").split(".")
    return min(len(parts), 5)


def segments_to_adoc(
    segments: list[dict],
    norm: str,
    version: str,
    source_path: Path | None = None,
) -> str:
    lines: list[str] = []
    src_label = source_path.name if source_path else "(source)"
    lines += [
        f"= {norm} {version}",
        ":doctype: article",
        ":source-highlighter: highlight.js",
        ":icons: font",
        ":toc: left",
        ":toclevels: 4",
        ":numbered:",
        "",
        "// Auto-generated by pdf-segment.py",
        f"// Source: {src_label}",
        "// Anchors link back to the original document",
        "",
    ]
    skip_types = {"HEADER", "FOOTER", "TOC", "OTHER"}
    for seg in segments:
        seg_type = seg["type"]
        if seg_type in skip_types:
            continue
        text    = seg["text"].strip()
        page_nr = seg["page"]
        anchor  = seg["id"]
        section = seg["section"]

        if seg_type == "SECTION":
            depth  = _section_depth(section)
            prefix = "=" * (depth + 1)
            anchor_label = f"{section} · p.{page_nr}" if section else f"p.{page_nr}"
            lines += ["", f"[[{anchor},{anchor_label}]]", f"{prefix} {text}", ""]

        elif seg_type in ("NORM", "INFORM"):
            role = "normative" if seg_type == "NORM" else "informative"
            kw   = seg.get("normative_keywords", [])
            kw_comment = f" // {', '.join(kw)}" if kw else ""
            lines += [
                f"[.{role}]#{kw_comment}",
                f"// {anchor} — para/page {page_nr}",
                text, "",
            ]

        elif seg_type == "TABLE":
            if "table_header" in seg:
                header = seg["table_header"]
                rows   = seg.get("table_rows", [])
                continued = seg.get("continues_on_next_page", False)
                merged_from = seg.get("merged_from_page")
                lines += ["", f"// {anchor} — page {page_nr}"]
                if continued:
                    lines.append("// ⚠️ table continues on next page")
                if merged_from:
                    lines.append(f"// ✅ merged with continuation from page {merged_from}")
                lines += ["|==="]
                lines.append(" | ".join(f"{c}" for c in header))
                for row in rows:
                    lines.append(" | ".join(f"{c}" for c in row))
                lines += ["|===", ""]
            else:
                lines += ["", f"// {anchor} — page {page_nr}",
                          f'[caption="{text[:80]}"]', "----", text, "----", ""]
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# File I/O
# ─────────────────────────────────────────────────────────────────────────────

def _write_segments(
    segments: list[dict],
    norm: str,
    version: str,
    stem: str,
    corpus_root: Path,
    source_path: Path | None,
    source_kind: str = "pdf",
) -> tuple[Path, Path]:
    seg_dir  = corpus_root / "specs" / "_segments"
    adoc_dir = corpus_root / "specs" / "_adoc"
    seg_dir.mkdir(parents=True, exist_ok=True)
    adoc_dir.mkdir(parents=True, exist_ok=True)

    seg_path  = seg_dir  / f"{stem}.segments.json"
    adoc_path = adoc_dir / f"{stem}.adoc"

    meta = {
        "norm":         norm,
        "version":      version,
        "source_kind":  source_kind,
        f"source_{source_kind}": str(source_path) if source_path else None,
        "segmented_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "segment_count": len(segments),
        "type_counts": {
            t: sum(1 for s in segments if s["type"] == t)
            for t in SEGMENT_TYPES
        },
        "segments": segments,
    }
    seg_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

    adoc_content = segments_to_adoc(segments, norm, version, source_path)
    adoc_path.write_text(adoc_content, encoding="utf-8")
    return seg_path, adoc_path


def iter_corpus_jsons(corpus_root: Path) -> Iterator[Path]:
    specs_dir = corpus_root / "specs"
    if not specs_dir.is_dir():
        return
    for p in sorted(specs_dir.glob("*.json")):
        yield p


def process_corpus_json(
    corpus_json: Path,
    profile: Profile,
    corpus_root: Path,
    force: bool = False,
    verbose: bool = True,
) -> str:
    try:
        rec = json.loads(corpus_json.read_text(encoding="utf-8"))
    except Exception as exc:
        print(f"[ERROR] Cannot read {corpus_json}: {exc}", file=sys.stderr)
        return "error"

    norm      = rec.get("norm",    corpus_json.stem)
    version   = rec.get("version", "")
    stem      = _safe_stem(corpus_json.stem)
    shortname = _format_shortname(rec)

    seg_path = corpus_root / "specs" / "_segments" / f"{stem}.segments.json"
    if not force and seg_path.exists():
        if verbose:
            print(f"[⏭️]  {corpus_json.name}  ({shortname})  — already segmented")
        return "skipped"

    source_path, source_kind = find_source_file(rec, stem)
    if source_path is None:
        if verbose:
            print(f"[⚠️]  {corpus_json.name}  ({shortname})  — no source file found")
        return "nosource"

    if verbose:
        icon = "📄" if source_kind == "docx" else "📋"
        print(f"[🔄]  {corpus_json.name}  ({shortname})  [{icon} {source_kind.upper()}]")

    try:
        if source_kind == "docx":
            segments = segment_docx(source_path, profile, doc_stem=stem)
        else:
            segments = segment_pdf(source_path, profile, doc_stem=stem)
    except Exception as exc:
        print(f"[ERROR] {corpus_json.name}: segmentation failed: {exc}", file=sys.stderr)
        return "error"

    seg_path, adoc_path = _write_segments(
        segments, norm, version, stem, corpus_root, source_path, source_kind
    )

    if verbose:
        counts = {t: sum(1 for s in segments if s["type"] == t) for t in SEGMENT_TYPES}
        print(
            f"   ✓  {len(segments)} blocks  |  "
            f"NORM={counts.get('NORM', 0)}  "
            f"SECTION={counts.get('SECTION', 0)}  "
            f"TABLE={counts.get('TABLE', 0)}  "
            f"INFORM={counts.get('INFORM', 0)}  "
            f"OTHER={counts.get('OTHER', 0)}"
        )
        print(f"   📄  {seg_path}")
        print(f"   📝  {adoc_path}")
    return "processed"


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="ETSI document → structural segments + AsciiDoc (DOCX preferred, PDF fallback)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python3 scripts/pdf-segment.py                          # all corpus specs
  python3 scripts/pdf-segment.py corpus/specs/ts_119403v020201p.json
  python3 scripts/pdf-segment.py --pdf downloads/specs/EN/en319401v020201p.pdf
  python3 scripts/pdf-segment.py --docx downloads/specs/EN/ESI-0019401v331v322.docx
  python3 scripts/pdf-segment.py --force                  # overwrite existing
  python3 scripts/pdf-segment.py --profile etsi-contribution --pdf ESI-0019478.pdf
  python3 scripts/pdf-segment.py --stats                  # summary only
  python3 scripts/pdf-segment.py --scan-styles downloads/specs/EN/ESI-0019401v331v322.docx
""",
    )
    parser.add_argument(
        "inputs", nargs="*",
        help="corpus/*.json files. If empty: all corpus/specs/*.json",
    )
    parser.add_argument(
        "--pdf", metavar="PDF_PATH",
        help="Process a raw PDF directly (bypasses corpus JSON)",
    )
    parser.add_argument(
        "--docx", metavar="DOCX_PATH",
        help="Process a raw DOCX directly (bypasses corpus JSON)",
    )
    parser.add_argument(
        "--scan-styles", metavar="DOCX_PATH",
        help="Print all paragraph style names found in a DOCX (style audit, no segmentation)",
    )
    parser.add_argument(
        "--profile", default="etsi-spec",
        choices=list(PROFILES),
        help="Segmentation profile (default: etsi-spec)",
    )
    parser.add_argument(
        "--corpus", default="corpus",
        help="Corpus root directory (default: corpus/)",
    )
    parser.add_argument("--force",  "-f", action="store_true", help="Overwrite existing")
    parser.add_argument("--quiet",  "-q", action="store_true", help="Suppress per-block output")
    parser.add_argument("--stats",        action="store_true", help="Print type statistics")
    args = parser.parse_args()

    if not HAS_PDFPLUMBER and not HAS_DOCX:
        sys.exit("[ERROR] Neither pdfplumber nor python-docx installed.\n"
                 "        Run: pip install pdfplumber python-docx")

    profile     = PROFILES[args.profile]
    corpus_root = Path(args.corpus)
    verbose     = not args.quiet

    # ── Style audit mode ──────────────────────────────────────────
    if args.scan_styles:
        import collections
        p = Path(args.scan_styles)
        if not p.is_file():
            sys.exit(f"[ERROR] File not found: {p}")
        doc = docx_lib.Document(str(p))
        styles: collections.Counter[str] = collections.Counter()
        samples: dict[str, str] = {}
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                sn = para.style.name or ""
                styles[sn] += 1
                if sn not in samples:
                    samples[sn] = t[:70]
        print(f"\nStyles in {p.name}:")
        for name, n in sorted(styles.items()):
            known_tag = "" if _is_known_style(name.lower()) else "  ← UNKNOWN"
            mapped    = _ETSI_STYLE_MAP.get(name.lower(), "")
            map_tag   = f"  → {mapped}" if mapped else ""
            print(f"  {n:4}x  {name:<30}{map_tag}{known_tag}")
            print(f"         {samples[name]!r}")
        return

    # ── Direct DOCX mode ──────────────────────────────────────────
    if args.docx:
        docx_path = Path(args.docx)
        if not docx_path.is_file():
            sys.exit(f"[ERROR] DOCX not found: {docx_path}")
        stem = _safe_stem(docx_path.stem)
        print(f"[🔄]  {docx_path.name}  (profile: {profile.name})  [📄 DOCX]")
        segments = segment_docx(docx_path, profile, doc_stem=stem)
        seg_path, adoc_path = _write_segments(
            segments, docx_path.stem, "", stem, corpus_root, docx_path, "docx"
        )
        if verbose:
            counts = {t: sum(1 for s in segments if s["type"] == t) for t in SEGMENT_TYPES}
            print(f"   ✓  {len(segments)} blocks total")
            for t in SEGMENT_TYPES:
                if counts.get(t, 0):
                    print(f"      {t:12} {counts[t]}")
            print(f"   📄  {seg_path}")
            print(f"   📝  {adoc_path}")
        return

    # ── Direct PDF mode ───────────────────────────────────────────
    if args.pdf:
        pdf_path = Path(args.pdf)
        if not pdf_path.is_file():
            sys.exit(f"[ERROR] PDF not found: {pdf_path}")
        stem = _safe_stem(pdf_path.stem)
        print(f"[🔄]  {pdf_path.name}  (profile: {profile.name})  [📋 PDF]")
        segments = segment_pdf(pdf_path, profile, doc_stem=stem)
        seg_path, adoc_path = _write_segments(
            segments, pdf_path.stem, "", stem, corpus_root, pdf_path, "pdf"
        )
        if verbose:
            counts = {t: sum(1 for s in segments if s["type"] == t) for t in SEGMENT_TYPES}
            print(f"   ✓  {len(segments)} blocks total")
            for t in SEGMENT_TYPES:
                if counts.get(t, 0):
                    print(f"      {t:12} {counts[t]}")
            print(f"   📄  {seg_path}")
            print(f"   📝  {adoc_path}")
        return

    # ── Corpus JSON mode ──────────────────────────────────────────
    if args.inputs:
        corpus_jsons = [Path(p) for p in args.inputs]
    else:
        corpus_jsons = list(iter_corpus_jsons(corpus_root))
        if not corpus_jsons:
            sys.exit(
                f"[ERROR] No corpus JSON files found in {corpus_root / 'specs'}.\n"
                "        Run: npm run ingest"
            )

    if verbose:
        docx_ok = "✅" if HAS_DOCX       else "❌ (pip install python-docx)"
        pdf_ok  = "✅" if HAS_PDFPLUMBER else "❌ (pip install pdfplumber)"
        print(f"[INFO]  {len(corpus_jsons)} corpus JSON(s)  "
              f"(profile: {profile.name}, "
              f"{'--force' if args.force else 'idempotent'})")
        print(f"        DOCX support: {docx_ok}")
        print(f"        PDF  support: {pdf_ok}\n")

    done = skipped_done = skipped_nosource = errors = 0
    total_segments = 0
    type_totals: dict[str, int] = {t: 0 for t in SEGMENT_TYPES}
    docx_count = pdf_count = 0

    for cj in corpus_jsons:
        try:
            result = process_corpus_json(
                cj, profile, corpus_root, force=args.force, verbose=verbose
            )
            if result == "processed":
                done += 1
                stem     = _safe_stem(cj.stem)
                seg_path = corpus_root / "specs" / "_segments" / f"{stem}.segments.json"
                if seg_path.exists():
                    try:
                        data = json.loads(seg_path.read_text(encoding="utf-8"))
                        total_segments += data.get("segment_count", 0)
                        kind = data.get("source_kind", "pdf")
                        if kind == "docx":
                            docx_count += 1
                        else:
                            pdf_count += 1
                        for t in SEGMENT_TYPES:
                            type_totals[t] += data.get("type_counts", {}).get(t, 0)
                    except Exception:
                        pass
            elif result == "skipped":
                skipped_done += 1
            elif result == "nosource":
                skipped_nosource += 1
            else:
                errors += 1
        except Exception as exc:
            errors += 1
            print(f"[ERROR] {cj.name}: {exc}", file=sys.stderr)

    if verbose or args.stats:
        print(f"\n{'─'*60}")
        print(f"  Processed   : {done}  (DOCX: {docx_count}, PDF: {pdf_count})")
        print(f"  Already done: {skipped_done}")
        print(f"  No source   : {skipped_nosource}")
        print(f"  Errors      : {errors}")
        print(f"  Segments    : {total_segments:,}")
        if args.stats and total_segments:
            print(f"\n  Type breakdown:")
            for t in SEGMENT_TYPES:
                n = type_totals.get(t, 0)
                if n:
                    pct = n / total_segments * 100
                    print(f"    {t:12} {n:6,}  ({pct:5.1f}%)")
        print(f"{'─'*60}")


if __name__ == "__main__":
    main()
